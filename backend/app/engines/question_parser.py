"""
Question Paper Parser Engine.
Parses question paper PDFs and DOCX files to extract questions, options, and media.

Data model:
  - 1 Question Paper PDF = 1 Medium + 1 Class + MULTIPLE Subjects
  - The PDF header contains subject-wise Q-number ranges.
  - Parser auto-detects Class, Medium from the header text.
  - Parser uses the curriculum registry to assign subjects based on Q-number ranges.
"""
import re
import os
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from app.models.schemas import ParsedQuestion, ParsedOption, QuestionType
from app.config import UPLOAD_DIR, MEDIA_DIR
from app.rules.curriculum_registry import get_subject_for_question, get_subject_ranges

logger = logging.getLogger(__name__)


class QuestionPaperParser:
    """
    Engine to parse question papers from PDF and DOCX formats.
    Auto-detects class, medium, and subject boundaries from the PDF header.
    Tags each question with its subject based on Q-number ranges.
    """

    QUESTION_PATTERNS = [
        re.compile(r'^\s*(?:Q\.?\s*)?(\d+)\s*[.)-]\s*(.+)', re.MULTILINE),
        re.compile(r'^\s*Question\s+(\d+)\s*[.:-]\s*(.+)', re.MULTILINE | re.IGNORECASE),
    ]

    OPTION_PATTERNS = [
        re.compile(r'^\s*\(?([A-Da-d])\s*[.)-]\)?\s*(.+)', re.MULTILINE),
    ]


    SECTION_PATTERNS = [
        re.compile(r'^\s*Section\s*[-:]?\s*([A-Z])\s*[-:]?\s*(.*)', re.MULTILINE | re.IGNORECASE),
        re.compile(r'^\s*Part\s*[-:]?\s*([A-Z])\s*[-:]?\s*(.*)', re.MULTILINE | re.IGNORECASE),
    ]

    # Patterns to detect class from header
    CLASS_PATTERNS = [
        re.compile(r'Class\s*[:\-]?\s*([IVXLCDM]+|\d+)', re.IGNORECASE),
        re.compile(r'Grade\s*[:\-]?\s*(\d+)', re.IGNORECASE),
    ]

    # Patterns to detect medium from header
    MEDIUM_PATTERNS = [
        re.compile(r'Medium\s*[:\-]?\s*([A-Za-z]+)', re.IGNORECASE),
        re.compile(r'([A-Za-z]+)\s+Medium', re.IGNORECASE),
    ]

    # Patterns to detect subject Q-ranges from header
    # e.g., "Language 1 (Assamese) – Q. No. 1 - 10" or "Mathematics: Q. No. 21 – 30"
    SUBJECT_RANGE_PATTERNS = [
        re.compile(
            r'([A-Za-z\s\(\)]+?)\s*[-–:]\s*Q\.?\s*(?:No\.?)?\s*(\d+)\s*[-–]\s*(\d+)',
            re.IGNORECASE
        ),
        re.compile(
            r'([A-Za-z\s\(\)]+?)\s*[-–]\s*(?:Q\.?\s*(?:No\.?)?\s*)?(\d+)\s*[-–]\s*(\d+)',
            re.IGNORECASE
        ),
    ]

    # Roman numeral conversion
    ROMAN_MAP = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
                 'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10}

    def __init__(self):
        self.current_section = None
        self.extracted_images = []
        self.media_map = {}

    def parse_file(self, file_path: str, medium: str = "",
                   class_level: int = 0) -> List[ParsedQuestion]:
        """
        Parse a question paper file (PDF or DOCX).

        Args:
            file_path: Path to the file
            medium: Language medium (optional, auto-detected from header)
            class_level: Class/grade level (optional, auto-detected from header)

        Returns:
            List of parsed questions, each tagged with auto-detected subject
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path, medium, class_level)
        elif ext == ".docx":
            return self._parse_docx(file_path, medium, class_level)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_pdf(self, file_path: Path, medium: str,
                   class_level: int) -> List[ParsedQuestion]:
        """Parse a PDF question paper using fitz (PyMuPDF) for better character mapping."""
        try:
            import fitz
        except ImportError:
            return self._parse_pdf_plumber(file_path, medium, class_level)

        questions = []
        full_text = ""

        try:
            doc = fitz.open(str(file_path))
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text") or ""
                full_text += page_text + "\n"
            doc.close()

            # Still use pdfplumber for tables if needed, or just extract images
            try:
                import pdfplumber
                with pdfplumber.open(str(file_path)) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables):
                                self._save_table_as_image(table, "GEN",
                                                          class_level, page_num, table_idx)
            except ImportError:
                pass

            questions = self._parse_text_content(full_text, medium, class_level)
            self._extract_images_pymupdf(file_path, "GEN", class_level)

        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise

        return questions

    def _parse_pdf_plumber(self, file_path: Path, medium: str,
                          class_level: int) -> List[ParsedQuestion]:
        """Fallback PDF parsing using pdfplumber."""
        import pdfplumber
        full_text = ""
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                full_text += (page.extract_text() or "") + "\n"
        return self._parse_text_content(full_text, medium, class_level)

    def _parse_pdf_fallback(self, file_path: Path, medium: str,
                            class_level: int) -> List[ParsedQuestion]:
        """Fallback PDF parsing using PyMuPDF."""
        try:
            import fitz
        except ImportError:
            raise ImportError("Neither pdfplumber nor PyMuPDF is installed")

        full_text = ""
        doc = fitz.open(str(file_path))
        for page in doc:
            full_text += page.get_text() + "\n"
        doc.close()

        return self._parse_text_content(full_text, medium, class_level)

    def _get_text_from_para(self, p) -> str:
        """Extract text from a docx Paragraph, handling superscripts and symbols."""
        text = ""
        for run in p.runs:
            rt = run.text
            # Handle superscript 0/o/O as degree symbol
            if run.font.superscript and rt.lower() in ['0', 'o']:
                text += "°"
            # Handle Symbol font mapping for angles and other math symbols
            elif run.font.name == 'Symbol':
                # Comprehensive Symbol font mapping
                symbol_map = {
                    '\xd0': '∠', '\xb0': '°', '\x3d': '=', '\x2b': '+', '\x2d': '-',
                    '\x3c': '<', '\x3e': '>', '\x2a': '×', '\x2f': '÷', '\xb1': '±',
                    '\x44': 'Δ', '\x70': 'π', '\x53': 'Σ', '\x4f': 'Ω', '\x61': 'α',
                    '\x62': 'β', '\x67': 'γ', '\x64': 'δ', '\x71': 'θ', '\x6c': 'λ',
                    '\x6d': 'μ', '\x72': 'ρ', '\x73': 'σ', '\x74': 'τ', '\x66': 'φ',
                    '\x77': 'ω', '\xac': '¬', '\xd8': '∅', '\xb6': '∂'
                }
                for char in rt:
                    text += symbol_map.get(char, char)
            else:
                text += rt
        return text

    def _parse_docx(self, file_path: Path, medium: str,
                    class_level: int) -> List[ParsedQuestion]:
        """Parse a DOCX question paper, preserving text, images, and table order."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is not installed")

        doc = Document(str(file_path))
        full_text = ""
        table_counter = 0
        image_counter = 0
        global_q_counter = 1

        # Use a faster mapping for paragraphs and tables
        para_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        # Extract text in order, including tables and inline images
        for block in doc.element.body:
            if block.tag.endswith('p'):
                p = para_map.get(block)
                if p:
                    p_text = self._get_text_from_para(p)
                    full_text += p_text + "\n"
                    # Extract images inline
                    r_ids = re.findall(r'r:embed="([^"]+)"', p._element.xml)
                    for r_id in r_ids:
                        rel = doc.part.rels.get(r_id)
                        if rel and "image" in rel.reltype:
                            image_counter += 1
                            image_data = rel.target_part.blob
                            ext = os.path.splitext(rel.target_ref)[1] or ".png"
                            img_filename = f"GEN_{class_level}_IMG_{image_counter:03d}{ext}"
                            img_path = MEDIA_DIR / "images" / img_filename
                            with open(img_path, "wb") as f:
                                f.write(image_data)
                            self.extracted_images.append(str(img_path))
                            self.media_map[f"IMAGE_{image_counter}"] = str(img_path)
                            full_text += f"\n[IMAGE_{image_counter}]\n"
            
            elif block.tag.endswith('tbl'):
                t = table_map.get(block)
                if t:
                    table_counter += 1
                    table_data = []
                    table_text = ""
                    for row in t.rows:
                        row_data = []
                        is_question_row = False
                        
                        if len(row.cells) >= 2:
                            t0 = row.cells[0].text.strip()
                            t1 = row.cells[1].text.strip()
                            if t1 and t0 != t1:
                                # Check for numbering in cell 0
                                has_num = any('<w:numId' in p._element.xml for p in row.cells[0].paragraphs)
                                if (t0 == "" and has_num) or (re.match(r'^(?:Q\.?)?\s*\d+\.?$', t0)):
                                    is_question_row = True

                        if is_question_row:
                            c1 = row.cells[1]
                            q_lines = []
                            opt_counter = 0
                            opt_labels = ['A', 'B', 'C', 'D']
                            
                            for p in c1.paragraphs:
                                p_text = self._get_text_from_para(p)
                                if not p_text: continue
                                
                                if '<w:numId' in p._element.xml and len(q_lines) > 0:
                                    label = opt_labels[opt_counter] if opt_counter < 4 else 'D'
                                    q_lines.append(f"{label}. {p_text}")
                                    opt_counter += 1
                                else:
                                    q_lines.append(p_text)
                                    
                            c1_text = "\n".join(q_lines)
                            # Only add prefix if not already present to satisfy "exactly get parsed" requirement
                            if not re.match(r'^\s*(?:Q\.?\s*)?\d+', c1_text):
                                q_text = f"{global_q_counter}. {c1_text}"
                            else:
                                q_text = c1_text
                            global_q_counter += 1
                            
                            row_data = [row.cells[0].text.strip(), q_text] + [c.text.strip() for c in row.cells[2:]]
                            table_text += q_text + "\n[ROW_END]\n"

                        else:
                            row_txt_list = []
                            for cell in row.cells:
                                c_text = cell.text.strip()
                                row_data.append(c_text)
                                row_txt_list.append(c_text)
                            table_text += "\t".join(row_txt_list) + "\n[ROW_END]\n"
                            
                        table_data.append(row_data)
                    
                    full_text += f"\n[TABLE_START_{table_counter}]\n{table_text}[TABLE_END]\n"
                    
                    # Render table as image
                    save_path = self._save_table_as_image(table_data, "GEN", class_level, 1, table_counter - 1)
                    if save_path:
                        self.media_map[f"TABLE_{table_counter}"] = save_path

        return self._parse_text_content(full_text, medium, class_level)

    def _detect_class_from_header(self, text: str) -> int:
        """Auto-detect class/grade level from the PDF header text."""
        for pattern in self.CLASS_PATTERNS:
            m = pattern.search(text)
            if m:
                val = m.group(1).strip()
                # Try Roman numeral first
                if val.upper() in self.ROMAN_MAP:
                    return self.ROMAN_MAP[val.upper()]
                # Try integer
                try:
                    return int(val)
                except ValueError:
                    pass
        return 0

    def _detect_medium_from_header(self, text: str) -> str:
        """Auto-detect medium from the PDF header text."""
        for pattern in self.MEDIUM_PATTERNS:
            m = pattern.search(text)
            if m:
                return m.group(1).strip()
        return ""

    def _detect_subject_ranges_from_header(self, text: str) -> List[Dict]:
        """
        Detect subject → Q-number ranges from the PDF header.
        e.g., "Language 1 (Assamese) – Q. No. 1 - 10" → {subject: "Language 1 (Assamese)", start: 1, end: 10}
        """
        ranges = []
        for pattern in self.SUBJECT_RANGE_PATTERNS:
            for m in pattern.finditer(text):
                subject_name = m.group(1).strip().rstrip('-–: ')
                # Skip if subject name looks like noise (too short or is a number)
                if len(subject_name) < 2 or subject_name.isdigit():
                    continue
                try:
                    start = int(m.group(2))
                    end = int(m.group(3))
                    if start > 0 and end >= start:
                        ranges.append({
                            "subject": subject_name,
                            "start": start,
                            "end": end,
                        })
                except (ValueError, IndexError):
                    continue

        # Deduplicate by range
        seen = set()
        unique = []
        for r in ranges:
            key = (r["start"], r["end"])
            if key not in seen:
                seen.add(key)
                unique.append(r)

        return sorted(unique, key=lambda x: x["start"])

    def _get_subject_for_qnum(self, q_num: int, detected_ranges: List[Dict],
                              class_level: int, medium: str) -> str:
        """
        Determine the subject for a question number.
        Priority: detected ranges from header > curriculum registry fallback.
        """
        # First try detected ranges from the PDF header
        for r in detected_ranges:
            if r["start"] <= q_num <= r["end"]:
                return r["subject"]

        # Fallback to curriculum registry
        from app.rules.curriculum_registry import get_subject_for_question
        subj = get_subject_for_question(q_num, class_level, medium)
        if subj != "Unknown":
            return subj

        return "General"

    def _parse_text_content(self, text: str, medium: str,
                            class_level: int) -> List[ParsedQuestion]:
        """
        Parse raw text content to extract questions and options.
        Auto-detects class, medium, and subject ranges from the header.
        """
        # ─── Auto-detect metadata from header (first ~500 chars) ──────
        header_text = text[:1500]  # Usually header info is in the first part

        detected_class = self._detect_class_from_header(header_text)
        detected_medium = self._detect_medium_from_header(header_text)
        detected_ranges = self._detect_subject_ranges_from_header(header_text)

        # Use detected values, fall back to provided values
        final_class = detected_class or class_level
        final_medium = detected_medium or medium

        if detected_class:
            logger.info(f"Auto-detected Class: {detected_class}")
        if detected_medium:
            logger.info(f"Auto-detected Medium: {detected_medium}")
        if detected_ranges:
            logger.info(f"Auto-detected subject ranges: {detected_ranges}")

        # If no ranges detected from header, use curriculum registry
        if not detected_ranges and final_class > 0:
            registry_ranges = get_subject_ranges(final_class, final_medium)
            if registry_ranges:
                detected_ranges = registry_ranges
                logger.info(f"Using curriculum registry ranges for Class {final_class}")

        # ─── Extract questions ────────────────────────────────────────
        questions = []
        current_section = None
        lines = text.split('\n')

        current_question = None
        current_options = []

        for line_num, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue

            # Check for section headers
            for pattern in self.SECTION_PATTERNS:
                sec_match = pattern.match(stripped)
                if sec_match:
                    current_section = sec_match.group(1)
                    break

            # Check for question patterns
            question_matched = False
            for pattern in self.QUESTION_PATTERNS:
                q_match = pattern.match(stripped)
                if q_match:
                    if current_question is not None:
                        current_question.options = current_options
                        questions.append(current_question)
                        current_options = []

                    q_num = int(q_match.group(1))
                    # Preserve exact text from document including number/prefix
                    q_text = stripped

                    # Auto-detect subject from Q number
                    subj = self._get_subject_for_qnum(
                        q_num, detected_ranges, final_class, final_medium
                    )

                    subject_code = subj.upper()[:3] if subj else "GEN"
                    # Clean up subject code for IDs (remove special chars)
                    subject_code = re.sub(r'[^A-Z0-9]', '', subject_code)[:3] or "GEN"
                    q_id = f"{subject_code}_{final_class}_Q{q_num}"

                    current_question = ParsedQuestion(
                        question_number=q_num,
                        question_id=q_id,
                        question_text=q_text,

                        subject=subj,
                        class_level=final_class,
                        medium=final_medium,
                        section=current_section,
                        raw_text=stripped,
                    )
                    question_matched = True
                    break

            if question_matched:
                continue

            # Check for media markers
            img_match = re.match(r'^\[(IMAGE_\d+)\]$', stripped)
            tbl_match = re.match(r'^\[(TABLE_START_\d+)\]$', stripped)
            
            if img_match and current_question is not None:
                media_key = img_match.group(1)
                if media_key in self.media_map:
                    current_question.media_file_path = self.media_map[media_key]
                    current_question.media_type = "image"
                continue

            if tbl_match and current_question is not None:
                media_key = tbl_match.group(1).replace("TABLE_START", "TABLE")
                if media_key in self.media_map:
                    current_question.media_file_path = self.media_map[media_key]
                    current_question.media_type = "image"
                continue
                
            # Skip TABLE_END and ROW_END markers to finalize the question
            if stripped in ["[TABLE_END]", "[ROW_END]"]:
                if current_question is not None:
                    current_question.options = current_options
                    questions.append(current_question)
                    current_question = None
                    current_options = []
                continue

            # Check for option patterns
            for pattern in self.OPTION_PATTERNS:
                o_match = pattern.match(stripped)
                if o_match and current_question is not None:
                    label = o_match.group(1).upper()
                    # Preserve exact text from document including label
                    text_val = stripped
                    current_options.append(ParsedOption(label=label, text=text_val))
                    break

            else:
                if current_question is not None:
                    cleaned_line = self._clean_text(stripped)
                    if current_options:
                        current_options[-1].text += "\n" + cleaned_line
                    else:
                        current_question.question_text += "\n" + cleaned_line

        # Save the last question
        if current_question is not None:
            current_question.options = current_options
            questions.append(current_question)

        # Determine question types
        for q in questions:
            if len(q.options) >= 2:
                q.question_type = QuestionType.MCQ
            elif any(opt.text.lower() in ["true", "false"] for opt in q.options):
                q.question_type = QuestionType.TRUE_FALSE
            else:
                q.question_type = QuestionType.SHORT_ANSWER

        # Log subject breakdown
        subject_counts = {}
        for q in questions:
            subject_counts[q.subject] = subject_counts.get(q.subject, 0) + 1
        logger.info(f"Parsed {len(questions)} questions — Subject breakdown: {subject_counts}")

        return questions

    def _clean_text(self, text: str) -> str:
        """Normalize math symbols and geometry characters."""
        if not text: return ""
        # Fix common degree mis-parses and normalize angle symbols
        text = text.replace('º', '°').replace('ª', '°').replace('^o', '°').replace('^0', '°')
        # Standardize angle symbol to U+2220
        text = text.replace('∡', '∠').replace('∢', '∠').replace('∟', '∠')
        # Fix common extraction artifacts
        text = text.replace('ï¿½', '°').replace('\u200b', '').replace('\ufeff', '')
        # Replace non-breaking spaces
        text = text.replace('\xa0', ' ')
        return text.strip()


    def _extract_images_pymupdf(self, file_path: Path, subject: str, class_level: int) -> List[str]:
        """Extract images from PDF using PyMuPDF for high quality extraction."""
        extracted = []
        try:
            import fitz
            doc = fitz.open(str(file_path))
            img_counter = 0

            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)

                for img_idx, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    if base_image:
                        img_counter += 1
                        img_ext = base_image.get("ext", "png")
                        subject_code = subject.upper()[:3] if subject else "GEN"
                        img_filename = f"{subject_code}_{class_level}_P{page_num + 1}_IMG_{img_counter}.{img_ext}"
                        img_path = MEDIA_DIR / "images" / img_filename
                        with open(img_path, "wb") as f:
                            f.write(base_image["image"])
                        extracted.append(str(img_path))
            doc.close()
        except ImportError:
            logger.warning("PyMuPDF not available for image extraction")
        except Exception as e:
            logger.error(f"Error extracting images: {e}")

        self.extracted_images.extend(extracted)
        return extracted

    def _save_table_as_image(self, table_data: list, subject: str, class_level: int,
                             page_num: int, table_idx: int) -> Optional[str]:
        """Convert a table to an image and save it."""
        try:
            from PIL import Image, ImageDraw, ImageFont

            if not table_data or not table_data[0]:
                return None

            rows = len(table_data)
            cols = len(table_data[0]) if table_data[0] else 0
            cell_width = 150
            cell_height = 30
            padding = 5

            img_width = cols * cell_width + 2 * padding
            img_height = rows * cell_height + 2 * padding

            img = Image.new('RGB', (img_width, img_height), 'white')
            draw = ImageDraw.Draw(img)

            for r_idx, row in enumerate(table_data):
                for c_idx, cell in enumerate(row or []):
                    x = padding + c_idx * cell_width
                    y = padding + r_idx * cell_height
                    draw.rectangle([x, y, x + cell_width, y + cell_height], outline='black')
                    cell_text = str(cell) if cell else ""
                    draw.text((x + 5, y + 5), cell_text[:20], fill='black')

            subject_code = subject.upper()[:3] if subject else "GEN"
            filename = f"{subject_code}_{class_level}_P{page_num}_TABLE_{table_idx + 1}.png"
            save_path = MEDIA_DIR / "tables" / filename
            img.save(str(save_path))
            return str(save_path)
        except Exception as e:
            logger.error(f"Error saving table as image: {e}")
            return None

    def get_extracted_images(self) -> List[str]:
        """Return list of all extracted image paths."""
        return self.extracted_images

    def reset(self):
        """Reset parser state."""
        self.current_section = None
        self.extracted_images = []
        self.media_map = {}
