"""
LO Mapping Parser Engine.
Parses Learning Outcome mapping PDFs and Excel files.

Data model:
  - 1 LO Mapping PDF = 1 Subject + 1 Class (medium-independent)
  - The PDF header contains Subject and Class info.
  - Inside: table with Class/Grade, Subject, Domain, Strand, Learning Outcome, Item ID,
    Cognitive Level, Answer Key
  - Item IDs correspond to question numbers in the question paper.
"""
import re
import logging
from pathlib import Path
from typing import List, Dict
from app.models.schemas import LOMapping

logger = logging.getLogger(__name__)


class LOMappingParser:
    HEADER_KEYWORDS = {
        "item_id": ["item id", "item no", "sl no", "serial", "q.no", "qno"],
        "domain": ["domain", "content domain"],
        "strand": ["strand", "sub-domain"],
        "learning_outcome": ["learning outcome", "lo", "outcome", "competency"],
        "cognitive_level": ["cognitive level", "cognitive", "bloom"],
        "answer_key": ["answer", "answer key", "correct answer", "key", "ans"],
        "subject": ["subject"],
        "class_level": ["class", "grade"],
    }

    # Patterns to detect subject and class from header text (before the table)
    SUBJECT_HEADER_PATTERNS = [
        re.compile(r'Subject\s*[:\-]?\s*(.+?)(?:\n|$)', re.IGNORECASE),
    ]
    CLASS_HEADER_PATTERNS = [
        re.compile(r'Class\s*[:\-]?\s*([IVXLCDM]+|\d+)', re.IGNORECASE),
        re.compile(r'Grade\s*[:\-]?\s*(\d+)', re.IGNORECASE),
    ]
    ROMAN_MAP = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
                 'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10}

    def parse_file(self, file_path: str, subject: str = "", class_level: int = 0) -> List[LOMapping]:
        """
        Parse an LO mapping file.
        Subject and class_level are auto-detected from the PDF header if not provided.
        """
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(Path(file_path), subject, class_level)
        elif ext == ".docx":
            return self._parse_docx(Path(file_path), subject, class_level)
        elif ext in [".xlsx", ".xls"]:
            return self._parse_excel(Path(file_path), subject, class_level)
        raise ValueError(f"Unsupported file type: {ext}")

    def _detect_subject_from_header(self, text: str) -> str:
        """Auto-detect subject from the LO mapping PDF header."""
        for pattern in self.SUBJECT_HEADER_PATTERNS:
            m = pattern.search(text)
            if m:
                return m.group(1).strip().rstrip(':').strip()
        return ""

    def _detect_class_from_header(self, text: str) -> int:
        """Auto-detect class from the LO mapping PDF header."""
        for pattern in self.CLASS_HEADER_PATTERNS:
            m = pattern.search(text)
            if m:
                val = m.group(1).strip()
                if val.upper() in self.ROMAN_MAP:
                    return self.ROMAN_MAP[val.upper()]
                try:
                    return int(val)
                except ValueError:
                    pass
        return 0

    def _parse_pdf(self, file_path: Path, subject: str, class_level: int) -> List[LOMapping]:
        mappings = []
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                # Extract header text from first page to detect subject/class
                first_page_text = pdf.pages[0].extract_text() or "" if pdf.pages else ""
                detected_subject = self._detect_subject_from_header(first_page_text)
                detected_class = self._detect_class_from_header(first_page_text)

                final_subject = subject or detected_subject
                final_class = class_level or detected_class

                if detected_subject:
                    logger.info(f"LO Mapping: Auto-detected Subject: {detected_subject}")
                if detected_class:
                    logger.info(f"LO Mapping: Auto-detected Class: {detected_class}")

                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            mappings.extend(self._parse_table(table, final_subject, final_class))
                    else:
                        text = page.extract_text() or ""
                        mappings.extend(self._parse_text(text, final_subject, final_class))
        except ImportError:
            import fitz
            doc = fitz.open(str(file_path))
            # Extract header from first page
            if len(doc) > 0:
                first_text = doc[0].get_text()
                detected_subject = self._detect_subject_from_header(first_text)
                detected_class = self._detect_class_from_header(first_text)
                final_subject = subject or detected_subject
                final_class = class_level or detected_class
            else:
                final_subject = subject
                final_class = class_level

            for page in doc:
                mappings.extend(self._parse_text(page.get_text(), final_subject, final_class))
            doc.close()

        # Ensure all mappings have subject and class set
        for m in mappings:
            if not m.subject and subject:
                m.subject = subject
            if not m.class_level and class_level:
                m.class_level = class_level

        return mappings

    def _parse_docx(self, file_path: Path, subject: str, class_level: int) -> List[LOMapping]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is not installed")
            
        doc = Document(str(file_path))
        mappings = []
        
        current_class = class_level
        current_subject = subject
        
        for block in doc.element.body:
            if block.tag.endswith('p'):
                for p in doc.paragraphs:
                    if p._element == block:
                        text = p.text
                        detected_class = self._detect_class_from_header(text)
                        if detected_class:
                            current_class = detected_class
                        detected_subj = self._detect_subject_from_header(text)
                        if detected_subj:
                            current_subject = detected_subj
                        break
            elif block.tag.endswith('tbl'):
                for t in doc.tables:
                    if t._element == block:
                        table_data = []
                        for row in t.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        
                        mappings.extend(self._parse_table(table_data, current_subject, current_class))
                        break
                        
        for m in mappings:
            if not m.subject and subject:
                m.subject = subject
            if not m.class_level and class_level:
                m.class_level = class_level
                
        return mappings

    def _parse_excel(self, file_path: Path, subject: str, class_level: int) -> List[LOMapping]:
        import pandas as pd
        mappings = []
        xls = pd.ExcelFile(str(file_path))
        for sheet in xls.sheet_names:
            df = pd.read_excel(str(file_path), sheet_name=sheet)
            if df.empty:
                continue
            col_map = self._detect_columns(df.columns.tolist())
            if not col_map.get("item_id"):
                continue
            for _, row in df.iterrows():
                try:
                    item_id_val = row.get(col_map.get("item_id", ""), "")
                    if pd.isna(item_id_val) or str(item_id_val).strip() == "":
                        continue
                    item_id = int(float(str(item_id_val)))
                    mappings.append(LOMapping(
                        item_id=item_id,
                        domain=self._s(row.get(col_map.get("domain", ""), "")),
                        strand=self._s(row.get(col_map.get("strand", ""), "")),
                        learning_outcome=self._s(row.get(col_map.get("learning_outcome", ""), "")),
                        cognitive_level=self._s(row.get(col_map.get("cognitive_level", ""), "")),
                        answer_key=self._s(row.get(col_map.get("answer_key", ""), "")),
                        subject=self._s(row.get(col_map.get("subject", ""), "")) or subject,
                        class_level=self._i(row.get(col_map.get("class_level", ""), 0)) or class_level,
                    ))
                except Exception as e:
                    logger.warning(f"Row parse error: {e}")
        return mappings

    def _parse_table(self, table: list, subject: str, class_level: int) -> List[LOMapping]:
        mappings = []
        if not table or len(table) < 2:
            return mappings
        headers = [str(h).strip().lower() if h else "" for h in table[0]]
        col_map = self._detect_columns(headers)
        if not col_map.get("item_id"):
            return mappings
        for row in table[1:]:
            try:
                if not row or all(not c or str(c).strip() == "" for c in row):
                    continue
                idx_map = {}
                for field, header_val in col_map.items():
                    if header_val in headers:
                        idx_map[field] = headers.index(header_val)
                iid = idx_map.get("item_id")
                if iid is None or iid >= len(row) or not row[iid]:
                    continue
                item_id = int(float(str(row[iid]).strip()))
                def g(f):
                    i = idx_map.get(f)
                    return str(row[i]).strip() if i is not None and i < len(row) and row[i] else ""
                mappings.append(LOMapping(
                    item_id=item_id, domain=g("domain"), strand=g("strand"),
                    learning_outcome=g("learning_outcome"), cognitive_level=g("cognitive_level"),
                    answer_key=g("answer_key"),
                    subject=g("subject") or subject,
                    class_level=self._i(g("class_level")) or class_level,
                ))
            except Exception as e:
                logger.warning(f"Table row error: {e}")
        return mappings

    def _parse_text(self, text: str, subject: str, class_level: int) -> List[LOMapping]:
        mappings = []
        pattern = re.compile(
            r'(?:Item\s*(?:ID\s*)?)?(\d+)\s*[|\t,]+\s*(?:Answer\s*(?:Key\s*)?)?[:=]?\s*([A-Da-d])\s*'
            r'[|\t,]+\s*(?:LO|Learning\s*Outcome)\s*[:=]?\s*(.+?)(?:\n|$)', re.IGNORECASE)
        for m in pattern.finditer(text):
            mappings.append(LOMapping(item_id=int(m.group(1)), answer_key=m.group(2).upper(),
                                      learning_outcome=m.group(3).strip(), subject=subject, class_level=class_level))
        return mappings

    def _detect_columns(self, headers: list) -> Dict[str, str]:
        col_map = {}
        headers_lower = [str(h).strip().lower() for h in headers]
        for field, keywords in self.HEADER_KEYWORDS.items():
            for header in headers_lower:
                if any(kw in header for kw in keywords):
                    col_map[field] = header
                    break
        return col_map

    @staticmethod
    def _s(val) -> str:
        if val is None: return ""
        try:
            import pandas as pd
            if isinstance(val, float) and pd.isna(val): return ""
        except: pass
        return str(val).strip()

    @staticmethod
    def _i(val) -> int:
        try: return int(float(str(val)))
        except: return 0
